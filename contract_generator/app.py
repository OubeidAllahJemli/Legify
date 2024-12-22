from flask import Flask, request, jsonify, render_template
import openai
from datetime import datetime
import os
from dotenv import load_dotenv
import json
from PyPDF2 import PdfReader
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from werkzeug.utils import secure_filename
from io import BytesIO
from PIL import Image
import pytesseract

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Configure OpenAI
openai.api_key = os.getenv('OPENAI_API_KEY')

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_image(image_file):
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        return str(e)
    
# Function to extract text from a PDF
def extract_pdf_text(pdf_file):
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return str(e)

# Function to split text into smaller chunks for embedding
def split_text_into_chunks(text, chunk_size=500):
    words = text.split()
    chunks = [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]
    return chunks

# Function to generate embeddings for each chunk
def generate_embeddings(text_chunks):
    """
    Generate embeddings for a list of text chunks using the updated OpenAI API.
    """
    try:
        response = openai.Embedding.create(
            model="text-embedding-ada-002",  # This is the recommended embedding model
            input=text_chunks
        )
        embeddings = [data["embedding"] for data in response["data"]]
        return embeddings
    except Exception as e:
        print(f"Error generating embeddings: {e}")
        return []


# Preprocess the PDF and create embeddings
pdf_path = "C:/Users/ghass/Downloads/Employment-law-overview-2022.pdf"

try:
    print("Extracting text from PDF...")
    legal_text = extract_pdf_text(pdf_path)
    print("Splitting text into chunks...")
    text_chunks = split_text_into_chunks(legal_text)
    print("Generating embeddings...")
    embeddings = generate_embeddings(text_chunks)
    print("Embeddings generated successfully.")
except Exception as e:
    print(f"An error occurred: {e}")

# Function to find the most relevant chunks based on query
def search_with_embeddings(query, text_chunks, embeddings):
    try:
        query_embedding = generate_embeddings([query])[0]
        similarities = cosine_similarity([query_embedding], embeddings)[0]
        top_indices = np.argsort(similarities)[::-1][:5]  # Get top 5 matches
        relevant_chunks = [text_chunks[i] for i in top_indices]
        return " ".join(relevant_chunks)
    except Exception as e:
        return f"Error during search: {e}"

# Function to generate legal contracts
def generate_contract(template_type, details):
    from docx import Document
    from docx.shared import Pt
    from io import BytesIO
    from flask import send_file

    try:
        # Mapping des types de contrats
        contract_types = {
            "CDI": "Contrat à Durée Indéterminée (Permanent Contract)",
            "CDD": "Contrat à Durée Déterminée (Fixed-Term Contract)",
            "Intérim": "Contrat de Travail Temporaire (Temporary Work Contract)",
            "Freelance": "Contrat de Prestataire Indépendant (Freelance Contract)"
        }

        # Création du prompt pour GPT
        prompt = f"""
        Generate a detailed legal contract for {contract_types.get(template_type, 'this employment type')} with these specifics:
        {json.dumps(details, indent=2)}
        Ensure the contract includes all standard legal clauses and specific terms provided.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a legal contract writer specializing in drafting legally binding contracts between an employer and an employee. Your task is to ensure that all key elements of a contract, such as job description, compensation, confidentiality clauses, duration, and termination conditions, are clearly stated and aligned with the laws applicable in the relevant jurisdiction. Your role involves generating contracts that are professional, legally compliant, and tailored to the specific needs of both parties."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        contract_content = response.choices[0].message['content']

        # Création du document Word
        doc = Document()
        doc.add_heading('Employment Contract', level=0)
        style = doc.styles['Normal']
        style.font.name='Arial'
        style.font.size=Pt(12)

        for line in contract_content.split('\n'):
            if line.strip():  # Évite les lignes vides
                doc.add_paragraph(line)

        # Sauvegarde dans un buffer
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Renvoi du fichier
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name="employment_contract.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return f"An error occurred during contract generation: {e}"


# Function to analyze contracts for risks and compliance
def analyze_contract(file_path, file_type, analysis_type):
    # Extract text from the file
    if file_type == 'pdf':
        contract_text = extract_pdf_text(file_path)
    elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
        contract_text = extract_text_from_image(file_path)
    else:
        return "Unsupported file type. Please provide a PDF or an image."

    # Check if text extraction was successful
    if not contract_text.strip():
        return "Failed to extract text from the file."

    # Split the extracted text into chunks and generate embeddings
    text_chunks = split_text_into_chunks(contract_text)
    embeddings = generate_embeddings(text_chunks)

    if not embeddings:
        return "Failed to generate embeddings."

    # Define prompts for different analysis types
    prompts = {
        "risk": "Analyze this contract for potential risks, ambiguous clauses, and missing terms:",
        "compliance": "Review this contract for compliance with standard regulations and highlight any issues:",
        "gdpr": "Specifically analyze this contract for GDPR compliance and data protection clauses:"
    }

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a highly skilled legal contract analyzer specializing in risk assessment, compliance, and ensuring adherence to relevant laws and regulations. Your role is to carefully review contract texts, identifying potential risks, compliance issues, and any legal concerns that may arise from ambiguous or unclear clauses. Based on the analysis type, you will assess the contract for risk factors, compliance with industry standards, and legal requirements, and provide insights or recommendations."},
                {"role": "user", "content": f"{prompts.get(analysis_type, prompts['risk'])}\n\n{contract_text}"}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        return str(e)

# Flask endpoints
@app.route('/api/legal/guidance', methods=['POST'])
def get_legal_guidance_endpoint():
    """
    Provide legal guidance based on user query.
    """
    try:
        data = request.get_json()
        if 'query' not in data:
            return jsonify({'error': 'Missing required field: query'}), 400
        
        query = data['query']
        relevant_text = search_with_embeddings(query, text_chunks, embeddings)
        
        if not relevant_text:
            return jsonify({'guidance': 'No relevant information found.'}), 404
        
        # Use GPT to generate the response based on relevant text
        prompt = f"Based on the following legal document, answer the query: {query}\n\n{relevant_text}"
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a legal guidance assistant with expertise in providing clear, concise, and actionable legal advice. Your role is to assist users by offering expert guidance on legal matters, explaining complex concepts in simple terms, and helping them navigate legal processes and obligations. Whether the user needs help with contract interpretation, compliance questions, or legal procedures, you provide well-reasoned, understandable answers tailored to their needs."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        return jsonify({'guidance': response['choices'][0]['message']['content']})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contracts/generate', methods=['POST'])
def create_contract():
    try:
        data = request.get_json()
        required_fields = ['template_type', 'details']
        if not all(field in data for field in required_fields):
            return jsonify({
                'error': 'Missing required fields',
                'required_fields': required_fields
            }), 400

        return generate_contract(data['template_type'], data['details'])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def upload_and_analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':  # Check if the file name is empty
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):  # Use the allowed_file function
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()

        # Read file content into memory
        file_bytes = BytesIO(file.read())
        
        if file_extension == 'pdf':
            # Extract text from PDF
            contract_text = extract_pdf_text(file_bytes)
        else:
            # Extract text from image
            contract_text = extract_text_from_image(file_bytes)
        
        # If no text extracted, return an error
        if not contract_text.strip():
            return jsonify({'error': 'Failed to extract text from the file'}), 400

        # Get the analysis type from the form data
        analysis_type = request.form.get('analysis_type', 'risk')

        # Define prompts for different analysis types
        prompts = {
            "risk": "Analyze this contract for potential risks, ambiguous clauses, and missing terms:",
            "compliance": "Review this contract for compliance with standard regulations and highlight any issues:",
            "gdpr": "Specifically analyze this contract for GDPR compliance and data protection clauses:"
        }

        try:
            # Send the extracted contract text to OpenAI for analysis
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a highly skilled legal contract analyzer specializing in risk assessment, compliance, and ensuring adherence to relevant laws and regulations. Your role is to carefully review contract texts, identifying potential risks, compliance issues, and any legal concerns that may arise from ambiguous or unclear clauses. Based on the analysis type, you will assess the contract for risk factors, compliance with industry standards, and legal requirements, and provide insights or recommendations."},
                    {"role": "user", "content": f"{prompts.get(analysis_type, prompts['risk'])}\n\n{contract_text}"}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            return jsonify({"analysis_result": response.choices[0].message.content})
        except Exception as e:
            return jsonify({'error': f'An error occurred during analysis: {e}'}), 500

    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/')
def index():
    return render_template('index.html')  # Serve the frontend interface

if __name__ == '__main__':
    app.run(debug=True)
